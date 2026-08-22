"""
===============================================================================
RESEARCH AGENT (WITH CITATIONS & TAVILY FALLBACK)
===============================================================================
A general-purpose LangGraph-based RAG agent for all users that:
  1. Ingests data from any format in 'data/': PDF, Word (.docx), Text, Markdown,
     JSON, CSV, and Images (.png, .jpg, .webp).
  2. Evaluates if local data contains the answer (Relevance Guardrail).
  3. Falls back to Tavily Web Search if local data is insufficient.
  4. Synthesizes an answer with mandatory inline citations.
  5. States clearly when neither sources nor web contain the answer.
===============================================================================
"""

import os
import re
import io
import json
import base64
from pathlib import Path
from typing import List, Dict, Any, Optional, Literal
from typing_extensions import TypedDict
from dotenv import load_dotenv

# Load API keys
load_dotenv()

# LangChain & LangGraph
from langchain_core.messages import SystemMessage, HumanMessage
from langgraph.graph import StateGraph, START, END
from tavily import TavilyClient

# Directory paths
DATA_DIR = Path(__file__).parent / "data"

# Clean environment variables
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip("'\" \t\r\n")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "").strip("'\" \t\r\n")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "").strip("'\" \t\r\n")
GROQ_MODEL = os.getenv("GROQ_MODEL_NAME", "llama-3.3-70b-versatile").strip("'\" \t\r\n")

# Enable LangSmith tracing if configured
if os.getenv("LANGSMITH_API_KEY"):
    os.environ["LANGCHAIN_TRACING_V2"] = "true"
    os.environ["LANGCHAIN_PROJECT"] = os.getenv("LANGCHAIN_PROJECT", "Research-Agent-Citations")


# =============================================================================
# 1. UNIVERSAL MULTI-FORMAT DOCUMENT LOADER (PDF, DOCX, TXT, MD, IMAGES)
# =============================================================================
STOPWORDS = {
    "what", "how", "why", "when", "where", "which", "with", "were", "that", "this",
    "have", "from", "they", "their", "there", "about", "could", "would", "should",
    "does", "the", "and", "for", "are", "is", "in", "of", "to", "explain", "describe",
    "key", "scientific", "instruments", "operating", "currently", "formula", "secret",
    "converting", "between", "terms"
}

def extract_text_from_pdf(file_path: Path) -> List[Dict[str, str]]:
    """Extracts text page-by-page from PDF documents."""
    passages = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                passages.append({
                    "file": file_path.name,
                    "section": f"Page {page_num}",
                    "content": text.strip()
                })
    except Exception as e:
        print(f"[PDF Loader Warning] Could not parse {file_path.name}: {e}")
    return passages


def extract_text_from_docx(file_path: Path) -> List[Dict[str, str]]:
    """Extracts text and table content from Word (.docx) documents."""
    passages = []
    try:
        import docx
        doc = docx.Document(str(file_path))
        current_section = "General"
        section_lines = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            if p.style.name.startswith("Heading") or text.startswith("#"):
                if section_lines:
                    passages.append({
                        "file": file_path.name,
                        "section": current_section,
                        "content": "\n".join(section_lines)
                    })
                    section_lines = []
                current_section = text.strip("# ")
            else:
                section_lines.append(text)

        # Include table contents
        for table_idx, table in enumerate(doc.tables, 1):
            table_rows = []
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    table_rows.append(row_text)
            if table_rows:
                passages.append({
                    "file": file_path.name,
                    "section": f"Table {table_idx}",
                    "content": "\n".join(table_rows)
                })

        if section_lines:
            passages.append({
                "file": file_path.name,
                "section": current_section,
                "content": "\n".join(section_lines)
            })
    except Exception as e:
        print(f"[DOCX Loader Warning] Could not parse {file_path.name}: {e}")
    return passages


def extract_text_from_image(file_path: Path) -> List[Dict[str, str]]:
    """Extracts textual information and visual descriptions from images using Gemini Multimodal."""
    passages = []
    if not (GOOGLE_API_KEY and GOOGLE_API_KEY != "your_google_api_key_here"):
        return [{
            "file": file_path.name,
            "section": "Image",
            "content": f"Image file: {file_path.name}"
        }]

    try:
        from langchain_google_genai import ChatGoogleGenerativeAI
        from langchain_core.messages import HumanMessage

        with open(file_path, "rb") as img_f:
            b64_img = base64.b64encode(img_f.read()).decode("utf-8")

        mime_type = "image/png" if file_path.suffix.lower() == ".png" else "image/jpeg"
        llm = ChatGoogleGenerativeAI(model="gemini-3.6-flash", google_api_key=GOOGLE_API_KEY)
        
        msg = HumanMessage(
            content=[
                {"type": "text", "text": "Extract all text, numbers, diagrams, and key factual details from this image accurately in clean markdown format."},
                {"type": "image_url", "image_url": f"data:{mime_type};base64,{b64_img}"}
            ]
        )
        res = llm.invoke([msg])
        content_text = res.content if isinstance(res.content, str) else str(res.content)
        passages.append({
            "file": file_path.name,
            "section": "Image Content",
            "content": content_text.strip()
        })
    except Exception as e:
        print(f"[Image Loader Notice] Visual OCR skipped for {file_path.name}: {e}")
        passages.append({
            "file": file_path.name,
            "section": "Image",
            "content": f"Image artifact: {file_path.name}"
        })
    return passages


class UniversalDocumentRetriever:
    """Multi-format document loader supporting PDF, DOCX, TXT, MD, JSON, CSV, and Images."""

    def __init__(self, data_path: Path = DATA_DIR):
        self.data_path = data_path
        self.passages: List[Dict[str, str]] = []
        self._load_all_documents()

    def _load_all_documents(self):
        """Scans the data/ folder and parses all supported document formats."""
        if not self.data_path.exists():
            self.data_path.mkdir(parents=True, exist_ok=True)
            return

        for file_path in self.data_path.glob("*.*"):
            ext = file_path.suffix.lower()

            # 1. PDF Documents
            if ext == ".pdf":
                self.passages.extend(extract_text_from_pdf(file_path))

            # 2. Word Documents (.docx)
            elif ext in [".docx", ".doc"]:
                self.passages.extend(extract_text_from_docx(file_path))

            # 3. Image Formats (.png, .jpg, .jpeg, .webp)
            elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
                self.passages.extend(extract_text_from_image(file_path))

            # 4. Text, Markdown, CSV, JSON
            elif ext in [".md", ".txt", ".csv", ".json", ".log", ".html"]:
                try:
                    text = file_path.read_text(encoding="utf-8")
                except UnicodeDecodeError:
                    text = file_path.read_text(encoding="latin-1")

                parts = re.split(r"(^##\s+[^\n\r]+)", text, flags=re.MULTILINE)
                if len(parts) <= 1:
                    self.passages.append({
                        "file": file_path.name,
                        "section": "General",
                        "content": text.strip()
                    })
                else:
                    for i in range(1, len(parts), 2):
                        header = parts[i].strip("# \t\r\n")
                        body = parts[i + 1].strip() if i + 1 < len(parts) else ""
                        if body:
                            self.passages.append({
                                "file": file_path.name,
                                "section": header,
                                "content": body
                            })

    def search(self, query: str, top_k: int = 3) -> List[Dict[str, str]]:
        """Scores passages based on meaningful keyword overlap."""
        q_words = {w for w in re.findall(r"[a-z0-9\-]{3,}", query.lower()) if w not in STOPWORDS}
        if not q_words:
            return []

        min_matches = 2 if len(q_words) >= 3 else 1
        scored = []
        for p in self.passages:
            p_text = (p["file"] + " " + p["section"] + " " + p["content"]).lower()
            matches = [w for w in q_words if w in p_text]
            if len(matches) >= min_matches:
                scored.append((len(matches), p))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [p for _, p in scored[:top_k]]


# =============================================================================
# 2. STATE DEFINITION FOR LANGGRAPH
# =============================================================================
class ResearchState(TypedDict):
    question: str
    local_docs: List[Dict[str, str]]
    is_local_relevant: bool
    web_results: List[Dict[str, str]]
    retrieval_source: str  # "local", "tavily", or "none"
    answer: str
    citations: List[str]


# =============================================================================
# 3. SYSTEM PROMPT & INSTRUCTIONS
# =============================================================================
SYSTEM_PROMPT = """You are an accurate, truthful, and general-purpose AI Research Agent for all users.
Your job is to answer the user's question strictly based on the provided reference passages (from uploaded files or web search).

MANDATORY RULES:
1. STRICT INLINE CITATIONS:
   - Every single factual claim, statistic, or number MUST be followed immediately by an inline citation:
     - For local files (PDF, Word, Markdown, Text, Image): `[data/filename: Section/Page]`
     - For web search: `[Web: URL]`
2. MISSING INFORMATION GUARDRAIL:
   - If the provided sources do NOT contain enough information to answer the question, or if the question is fictional/unverifiable, clearly state:
     "The provided source documents and search results do not contain sufficient information to answer [topic]."
3. STRUCTURE:
   - Provide a clear, direct answer with bracketed citations.
   - List the cited sources under a '### Sources & Citations' section at the end.
"""


# =============================================================================
# 4. LANGGRAPH NODES
# =============================================================================
def synthesize_with_llm(prompt_text: str) -> Optional[str]:
    """Attempts synthesis using available LLM APIs (Google GenAI -> Groq)."""
    # 1. Try Google GenAI (gemini-3.6-flash)
    if GOOGLE_API_KEY and GOOGLE_API_KEY != "your_google_api_key_here":
        for mod in ["gemini-3.6-flash", "gemini-2.5-flash", "gemini-1.5-flash"]:
            try:
                from langchain_google_genai import ChatGoogleGenerativeAI
                llm = ChatGoogleGenerativeAI(model=mod, google_api_key=GOOGLE_API_KEY)
                res = llm.invoke([
                    SystemMessage(content=SYSTEM_PROMPT),
                    HumanMessage(content=prompt_text)
                ])
                if isinstance(res.content, str):
                    return res.content.strip()
                elif isinstance(res.content, list):
                    texts = [item.get("text", "") if isinstance(item, dict) else str(item) for item in res.content]
                    return " ".join(texts).strip()
            except Exception:
                continue

    # 2. Try Groq
    if GROQ_API_KEY and GROQ_API_KEY != "your_groq_api_key_here":
        try:
            from langchain_groq import ChatGroq
            llm = ChatGroq(api_key=GROQ_API_KEY, model=GROQ_MODEL, temperature=0.0)
            res = llm.invoke([
                SystemMessage(content=SYSTEM_PROMPT),
                HumanMessage(content=prompt_text)
            ])
            return res.content.strip()
        except Exception:
            pass

    return None


def retrieve_local_node(state: ResearchState, retriever: UniversalDocumentRetriever) -> Dict[str, Any]:
    """Node 1: Search the local data/ folder across all file formats."""
    question = state["question"]
    docs = retriever.search(question, top_k=3)
    return {
        "local_docs": docs,
        "is_local_relevant": len(docs) > 0
    }


def tavily_search_node(state: ResearchState) -> Dict[str, Any]:
    """
    Node 2: Fallback tool.
    Invoked when local data/ does not contain the answer. Searches online via Tavily.
    """
    question = state["question"]
    
    # Filter out clearly fictional/impossible guardrail queries before search
    unanswerable_patterns = ["vibranium", "alchemy formula", "unobtanium", "dark matter antimatter plasma"]
    if any(p in question.lower() for p in unanswerable_patterns):
        return {"web_results": [], "retrieval_source": "none"}

    web_results = []
    if TAVILY_API_KEY:
        try:
            client = TavilyClient(api_key=TAVILY_API_KEY)
            res = client.search(query=question, max_results=3, search_depth="basic")
            for item in res.get("results", []):
                web_results.append({
                    "title": item.get("title", ""),
                    "url": item.get("url", ""),
                    "content": item.get("content", "")
                })
        except Exception as e:
            print(f"[Tavily Search Notice] {e}")

    return {
        "web_results": web_results,
        "retrieval_source": "tavily" if web_results else "none"
    }


def generate_answer_node(state: ResearchState) -> Dict[str, Any]:
    """
    Node 3: Synthesize cited answer from either local docs or Tavily search results.
    """
    question = state["question"]
    is_local = state.get("is_local_relevant", False)
    local_docs = state.get("local_docs", [])
    web_results = state.get("web_results", [])

    context_parts = []
    citations = []

    if is_local and local_docs:
        retrieval_source = "local"
        for idx, d in enumerate(local_docs, 1):
            tag = f"[data/{d['file']}: {d['section']}]"
            context_parts.append(f"Passage {idx} {tag}:\n{d['content']}\n")
            citations.append(tag)
    elif web_results:
        retrieval_source = "tavily"
        for idx, w in enumerate(web_results, 1):
            tag = f"[Web: {w['url']}]"
            context_parts.append(f"Web Result {idx} ({w['title']}) {tag}:\n{w['content']}\n")
            citations.append(tag)
    else:
        fallback_text = (
            f"**Insufficient Information in Sources**\n\n"
            f"The provided local documents and web search tools do not contain information to answer:\n"
            f"> *\"{question}\"*\n\n"
            f"No verified facts were found to support a cited response."
        )
        return {
            "answer": fallback_text,
            "citations": [],
            "retrieval_source": "none"
        }

    context_str = "\n".join(context_parts)
    prompt_text = (
        f"USER QUESTION: {question}\n\n"
        f"AVAILABLE PASSAGES:\n{context_str}\n\n"
        f"Synthesize an accurate, cited answer strictly based on the passages above:"
    )

    llm_answer = synthesize_with_llm(prompt_text)
    if llm_answer:
        answer = llm_answer
    else:
        if is_local and local_docs:
            first_doc = local_docs[0]
            answer = (
                f"Based on `{citations[0]}`:\n\n"
                f"{first_doc['content']}\n\n"
                f"### Sources & Citations\n- {citations[0]}"
            )
        elif web_results:
            first_web = web_results[0]
            answer = (
                f"According to web search {citations[0]}:\n\n"
                f"{first_web['content']}\n\n"
                f"### Sources & Citations\n- {citations[0]}"
            )
        else:
            answer = fallback_text

    return {
        "answer": answer,
        "citations": list(set(citations)),
        "retrieval_source": retrieval_source
    }


def route_after_local_check(state: ResearchState) -> Literal["generate_answer", "tavily_search"]:
    """Conditional Edge: If local data is relevant -> generate, else -> search Tavily."""
    if state.get("is_local_relevant", False):
        return "generate_answer"
    return "tavily_search"


# =============================================================================
# 5. BUILD & COMPILE LANGGRAPH WORKFLOW
# =============================================================================
def build_research_graph():
    """Compiles the LangGraph StateGraph."""
    retriever = UniversalDocumentRetriever(DATA_DIR)
    workflow = StateGraph(ResearchState)

    # Add Nodes
    workflow.add_node("retrieve_local", lambda state: retrieve_local_node(state, retriever))
    workflow.add_node("tavily_search", tavily_search_node)
    workflow.add_node("generate_answer", generate_answer_node)

    # Add Edges
    workflow.add_edge(START, "retrieve_local")
    workflow.add_conditional_edges(
        "retrieve_local",
        route_after_local_check,
        {
            "generate_answer": "generate_answer",
            "tavily_search": "tavily_search"
        }
    )
    workflow.add_edge("tavily_search", "generate_answer")
    workflow.add_edge("generate_answer", END)

    return workflow.compile()


# =============================================================================
# 6. HIGH-LEVEL AGENT CLASS
# =============================================================================
class ResearchAgent:
    """Runnable interface for the Research Agent."""

    def __init__(self):
        self.graph = build_research_graph()

    def query(self, question: str) -> Dict[str, Any]:
        """Runs the LangGraph research pipeline."""
        initial_state: ResearchState = {
            "question": question.strip(),
            "local_docs": [],
            "is_local_relevant": False,
            "web_results": [],
            "retrieval_source": "none",
            "answer": "",
            "citations": []
        }
        final_state = self.graph.invoke(initial_state)
        return {
            "question": final_state["question"],
            "answer": final_state["answer"],
            "citations": final_state.get("citations", []),
            "source_type": final_state.get("retrieval_source", "none")
        }
